#### imports 
import docx
import zipfile
from lxml import etree
from docx_utils.flatten import opc_to_flat_opc
from docx import Document
import json
import docx2txt,json, collections

#########  function defination area 
def readwordfile():
    global fname 
    print("1 : " , fname)
    if(fname == ""):
        return
    
    doc = docx.Document(fname)
    for i in doc.paragraphs:
        print(i.text)

def generatewordtextfile():
    global fname 
    print("1 : " , fname)
    if(fname == ""):
        return
    file2 = open(r"C:\Users\91922\Desktop\RAJ\code\outputfile\f1.txt","w+")
    doc = docx.Document(fname)
    for i in doc.paragraphs:
        file2.write(i.text)

    file2.close()

    print("\n\n Text file create in output directory . . . ")  



def callwordprocess():
    ##D:\project 2022-23\Resume key value pair extration\code\tempfolder\TYCOA31.docx
    global fname
    fname = input("Enter File name and path : ")
    readwordfile()





# def get_word_xml(docx_filename):
#    with open(fname) as f:
#       zip = zipfile.ZipFile(f)
#       xml_content = zip.read('outputfile/document.xml')
#    return xml_content



# def get_xml_tree(xml_string):
#    return etree.fromstring(fname)

def docx_to_xml():
    opc_to_flat_opc(fname, "outputfile/sample.xml")

# def docx_to_json():
    
#     document = Document(fname)
#     lines = [para.text for para in document.paragraphs]
#     print(lines)
#     lines = [line.partition('.') for line in lines]
#     print(lines)
#     # lines = [(int(row_num), row_text) for row_num, _, row_text in lines]
   
#     # lines = [(n, [txt.partition(':') for txt in row_text.split(',')]) for n, row_text in lines]
#     lines = { {key.strip(): val.strip() for key, _, val in row} for  row in lines}
#     json_result = json.dumps(lines)

def word_to_json_logic():
    # step 1 get docx text
    text = docx2txt.process(fname)
    # convert to list
    li = [x for x in text.split('\n')]
    li = list(filter(None, li))
    # print(li)
    # print("=========================================")

    x = li 
    print("{")
    print(x[0:4])

    subs = '}{'
    print(subs)
    #res = li[4:18]
    res = li[4] 
    print("{",li[4],":",li[11],",")
    print(li[5],":",li[12],",")
    print(li[6],":",li[13],",")
    print(li[7],":",li[14],",")
    print(li[8],":",li[15],",")
    print(li[9],":",li[16],",")
    print(li[10],":",li[17],"}")
    
    # subs = '}Program/Degree{'
    # print(subs)
    res = li[18]


    subs = ':'
    print("{", str(res),subs)
    res = li[21:24]
    #print( str(res))
    print(li[19],":",li[20:21])
    print(",",li[21],":",li[22:23])

    subs = '}\nProjects:'
    res = li[25:30]
    print(subs, str(res))
    # json list
    # json_li = []
    # # convert and store all values
    # for x in li:
    











    #     x = x[2:] # remove 1. 2. 3. ...
    #     y = x.split(',')
    #     print(y)
    #     d = collections.defaultdict()
    #     for m in y:
    #         z = m.split(':')
    #         z1 = [x.strip() for x in z]
    #         # d[z1[0]] = z1[1]
    #     json_li.append(z)
    # # JSON conversion
    
    print("}")
    #ok print(json.dumps(li, indent=4))    